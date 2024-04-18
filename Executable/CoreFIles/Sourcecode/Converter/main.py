import pandas as pd
import docx
import sys

def main ():
    target_path = sys.argv[1]
    input = sys.argv[2]
    df = pd.read_excel(input)

    for enum, futurefile in df.iterrows():
        metadata = []; content = ""
        for prop, val in futurefile.items():
            if prop != "text":
                metadata.append (prop + ":" + val)
            else: 
                content = val
        doc = docx.Document()
        metastr = "|".join(metadata)
        doc.add_paragraph("FileNr:" +  str(enum+1) +"|"+metastr)
        doc.add_paragraph("#######DONT CHANGE THE ABOVE############")
        doc.add_paragraph(content)
        doc.save(target_path + "\\" + "File" + str(enum+1) + ".docx")